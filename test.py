from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

def set_default_font(doc, font_name="Aptos", font_size=11):
    """
    Imposta il font predefinito per il documento modificando lo stile 'Normal'.

    :param doc: Oggetto Document di python-docx
    :param font_name: Nome del font (default: "Aptos")
    :param font_size: Dimensione del font in pt (default: 11)
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)


def add_question_to_doc(doc, question, options, correct_index):
    """
    Aggiunge una domanda e le sue opzioni a un documento DOCX esistente.

    :param doc: Oggetto Document di python-docx
    :param question: Testo della domanda
    :param options: Lista di opzioni
    :param correct_index: Indice della risposta corretta (0-based)
    """
    # Aggiunge la domanda in grassetto
    question_paragraph = doc.add_paragraph()
    question_run = question_paragraph.add_run(question)
    question_run.bold = False

    # Aggiunge le opzioni con un quadratino (☐) davanti
    for i, option in enumerate(options):
        #option_paragraph = doc.add_paragraph("☐ ", style="ListBullet")
        #option_run = option_paragraph.add_run(option)
        option_paragraph = doc.add_paragraph()
        option_paragraph.paragraph_format.space_after = 0  # Rimuove spazio extra dopo
        option_paragraph.paragraph_format.space_before = 0  # Rimuove spazio extra prima
        option_run = option_paragraph.add_run(f"☐ {option}")

        # Se è la risposta corretta, evidenzia in giallo
        if i == correct_index:
            option_run.font.highlight_color = 7  # Giallo

def create_quiz_docx(questions_data, file_path):
    """
    Crea un file DOCX e aggiunge una lista di domande con le loro opzioni.

    :param questions_data: Lista di tuple (domanda, lista opzioni, indice risposta corretta)
    :param file_path: Percorso di salvataggio del file DOCX
    """
    doc = Document()
    set_default_font(doc, font_name="Aptos", font_size=11)

    for idx, (question, options, correct_index) in enumerate(questions_data, start=1):
        add_question_to_doc(doc, f"{idx}. {question}", options, correct_index)
        doc.add_paragraph()  # Aggiunge una riga vuota tra le domande

    doc.save(file_path)
    print(f"File DOCX salvato come {file_path}")

# Esempio di utilizzo
if __name__ == "__main__":
    questions = [
        ("Qual è la capitale della Francia?", ["Roma", "Parigi", "Berlino", "Madrid"], 1),
        ("Qual è il pianeta più vicino al Sole?", ["Terra", "Marte", "Mercurio", "Giove"], 2),
        ("Chi ha scritto 'La Divina Commedia'?", ["Petrarca", "Boccaccio", "Dante", "Manzoni"], 2),
    ]

    create_quiz_docx(questions, "quiz.docx")

